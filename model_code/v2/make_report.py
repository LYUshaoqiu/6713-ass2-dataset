"""
Generate bert_report.docx for COMP6713 Group Project
BERT-Based Hate Speech Detection report covering:
  1. Zero-shot evaluation
  2. Fine-tuning on HateXPlain / TweetEval
  3. Cascaded classification (failed attempt)
  4. Domain adaptation (final solution)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Inches as In_, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

IMG_DIR = r'C:\Users\chenf\Desktop\6713\group\report'

doc = Document()

# ── Page margins ─────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.18)
    section.right_margin  = Inches(1.18)

# ── Default styles ────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# ── Helper: set font for a run / paragraph ────────────────────
def fmt(run, bold=False, italic=False, size=11, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    fmt(run, bold=True, size=13 if level == 1 else 11)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(13)
    for run in p.runs:
        fmt(run)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    fmt(run)
    return p

def caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    fmt(run, italic=True, size=10)
    return p

# ── Table helper ──────────────────────────────────────────────
HDR_BG  = '1F497D'   # dark blue header
ROW1_BG = 'DCE6F1'  # light blue alternating row
ROW2_BG = 'FFFFFF'  # white

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'BFBFBF')
        borders.append(el)
    tcPr.append(borders)

def make_table(doc, headers, rows, col_widths_cm):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'

    # Header row
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = Cm(col_widths_cm[i])
        set_cell_bg(cell, HDR_BG)
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        fmt(run, bold=True, size=10, color=(255,255,255))

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = t.rows[r_idx + 1]
        bg  = ROW1_BG if r_idx % 2 == 0 else ROW2_BG
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = Cm(col_widths_cm[c_idx])
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            fmt(run, size=10)
    return t

def add_inline_bold(doc, parts, space_after=6):
    """parts = list of (text, bold) tuples"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(13)
    for text, bold in parts:
        run = p.add_run(text)
        fmt(run, bold=bold)
    return p

def insert_two_images(doc, path1, cap1, path2, cap2, w=2.8):
    """Insert two images side by side using a 2-col borderless table."""
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (path, _) in enumerate([(path1, cap1), (path2, cap2)]):
        cell = t.rows[0].cells[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.isfile(path):
            run = p.add_run()
            run.add_picture(path, width=In_(w))
    for ci, (_, cap) in enumerate([(path1, cap1), (path2, cap2)]):
        cell = t.rows[1].cells[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cap)
        fmt(run, italic=True, size=9, color=(0x60, 0x60, 0x60))
    for row in t.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for side in ('top','left','bottom','right','insideH','insideV'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                borders.append(el)
            tcPr.append(borders)

# ═══════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(6)
title_p.paragraph_format.space_after  = Pt(4)
run = title_p.add_run('BERT-Based Hate Speech Detection:\nZero-Shot, Fine-Tuning, and Cascaded Classification')
fmt(run, bold=True, size=16, color=(0x1F, 0x49, 0x7D))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(16)
run = sub_p.add_run('COMP6713 Group Project — Model Section Report')
fmt(run, italic=True, size=11, color=(0x40, 0x40, 0x40))

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════════
# SECTION 1 — ZERO-SHOT
# ═══════════════════════════════════════════════════════════════
heading(doc, '1.  Zero-Shot Evaluation of Pre-trained Models  (5 Credits)', level=1)

heading(doc, '1.1  Approach', level=2)
body(doc,
     'Four pre-trained models are evaluated in a zero-shot setting without any task-specific '
     'fine-tuning. For each model, [CLS] token embeddings are extracted and compared via '
     'cosine similarity against three fixed label prototype embeddings '
     '("This is a normal, neutral, or positive comment." / '
     '"This is an offensive or rude comment using bad language." / '
     '"This is hate speech targeting a person or specific group."). '
     'No gradient updates are performed; each model is used directly out of the box.')

body(doc, 'The four models evaluated are:')
bullet(doc, 'BERT-base-uncased — general-purpose encoder, no hate speech knowledge')
bullet(doc, 'HateBERT (GroNLP/hateBERT) — BERT further pre-trained on Reddit abusive content')
bullet(doc, 'Twitter-RoBERTa-base-hate (cardiffnlp/twitter-roberta-base-hate) — RoBERTa pre-trained on tweets and fine-tuned on TweetEval hate detection')
bullet(doc, 'BERTweet (vinai/bertweet-base) — RoBERTa pre-trained on 850M English tweets')

heading(doc, '1.2  Results — HateXPlain Test Set', level=2)
caption(doc, 'Table 1.  Zero-Shot Performance on HateXPlain Test Set (n = 2,012)')
make_table(doc,
    headers=['Model', 'Accuracy', 'Macro F1', 'F1 Normal', 'F1 Offensive', 'F1 Hate'],
    rows=[
        ['BERT-base-uncased',        '0.3857', '0.2845', '0.5450', '0.1075', '0.2009'],
        ['HateBERT',                 '0.3429', '0.3367', '0.3693', '0.3683', '0.2725'],
        ['Twitter-RoBERTa-base-hate','0.3454', '0.3307', '0.2547', '0.4265', '0.3109'],
        ['BERTweet',                 '0.3131', '0.2538', '0.4082', '0.3262', '0.0269'],
    ],
    col_widths_cm=[5.0, 2.0, 2.0, 2.2, 2.8, 2.0]
)

# HateXPlain confusion matrices — 2×2 grid
IMG_DIR_ZS = r'C:\Users\chenf\Desktop\6713\group\report'
insert_two_images(doc,
    os.path.join(IMG_DIR_ZS, 'zeroshot_BERT-base-uncased_HateXPlain_cm.png'),
    'Figure 1a. BERT-base-uncased',
    os.path.join(IMG_DIR_ZS, 'zeroshot_HateBERT_HateXPlain_cm.png'),
    'Figure 1b. HateBERT',
    w=2.8)
insert_two_images(doc,
    os.path.join(IMG_DIR_ZS, 'zeroshot_Twitter-RoBERTa-base-hate_HateXPlain_cm.png'),
    'Figure 1c. Twitter-RoBERTa-base-hate',
    os.path.join(IMG_DIR_ZS, 'zeroshot_BERTweet_HateXPlain_cm.png'),
    'Figure 1d. BERTweet',
    w=2.8)
cap_hx = doc.add_paragraph()
cap_hx.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_hx.paragraph_format.space_after = Pt(8)
cap_run = cap_hx.add_run('Figure 1. Zero-shot confusion matrices on HateXPlain test set')
fmt(cap_run, italic=True, size=9, color=(0x60, 0x60, 0x60))

doc.add_paragraph()

heading(doc, '1.3  Results — TweetEval Test Set', level=2)
caption(doc, 'Table 2.  Zero-Shot Performance on TweetEval Test Set (n = 3,830)')
make_table(doc,
    headers=['Model', 'Accuracy', 'Macro F1', 'F1 Normal', 'F1 Offensive', 'F1 Hate'],
    rows=[
        ['BERT-base-uncased',        '0.5843', '0.2842', '0.7482', '0.0269', '0.0776'],
        ['HateBERT',                 '0.4778', '0.3421', '0.6816', '0.0836', '0.2612'],
        ['Twitter-RoBERTa-base-hate','0.1922', '0.1983', '0.2590', '0.0894', '0.2465'],
        ['BERTweet',                 '0.3251', '0.2477', '0.5106', '0.1117', '0.1207'],
    ],
    col_widths_cm=[5.0, 2.0, 2.0, 2.2, 2.8, 2.0]
)

insert_two_images(doc,
    os.path.join(IMG_DIR_ZS, 'zeroshot_BERT-base-uncased_TweetEval_cm.png'),
    'Figure 2a. BERT-base-uncased',
    os.path.join(IMG_DIR_ZS, 'zeroshot_HateBERT_TweetEval_cm.png'),
    'Figure 2b. HateBERT',
    w=2.8)
insert_two_images(doc,
    os.path.join(IMG_DIR_ZS, 'zeroshot_Twitter-RoBERTa-base-hate_TweetEval_cm.png'),
    'Figure 2c. Twitter-RoBERTa-base-hate',
    os.path.join(IMG_DIR_ZS, 'zeroshot_BERTweet_TweetEval_cm.png'),
    'Figure 2d. BERTweet',
    w=2.8)
cap_te = doc.add_paragraph()
cap_te.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_te.paragraph_format.space_after = Pt(8)
cap_run2 = cap_te.add_run('Figure 2. Zero-shot confusion matrices on TweetEval test set')
fmt(cap_run2, italic=True, size=9, color=(0x60, 0x60, 0x60))

doc.add_paragraph()

heading(doc, '1.4  Analysis', level=2)
body(doc,
     'All four models achieve modest Macro F1 scores (0.20–0.34) without fine-tuning, '
     'confirming that zero-shot CLS cosine similarity is insufficient for reliable '
     '3-class hate speech detection. Nevertheless, the results reveal meaningful '
     'differences across model types:')

add_inline_bold(doc, [
    ('HateBERT achieves the best overall balance. ', True),
    ('Macro F1 = 0.337 on HateXPlain and 0.342 on TweetEval — the highest across both '
     'datasets. Pre-training on Reddit abusive content gives it broader coverage of '
     'offensive language patterns, producing more balanced F1 scores across all three '
     'classes compared to the other models.', False)
])

add_inline_bold(doc, [
    ('Twitter-RoBERTa shows class-level specialisation. ', True),
    ('It achieves the highest F1 Offensive on HateXPlain (0.427) and the highest '
     'F1 Hate on TweetEval (0.247), consistent with its pre-training on hate-labelled '
     'tweet data. However, it dramatically over-predicts Offensive on TweetEval '
     '(Accuracy = 0.192), collapsing the Normal class F1 to 0.259.', False)
])

add_inline_bold(doc, [
    ('BERT-base-uncased defaults to predicting Normal. ', True),
    ('It achieves the highest F1 Normal on both datasets (0.545 / 0.748) but near-zero '
     'F1 for Offensive and Hate, indicating a strong majority-class bias in the absence '
     'of task-specific training.', False)
])

add_inline_bold(doc, [
    ('BERTweet performs weakest overall. ', True),
    ('Despite being pre-trained on 850M tweets, tweet-specific language understanding '
     'does not directly transfer to hate speech classification without fine-tuning. '
     'F1 Hate is near zero on HateXPlain (0.027), suggesting the model embeds hate '
     'speech in a space close to Normal content.', False)
])

body(doc,
     'These results establish a meaningful zero-shot baseline and demonstrate that '
     'domain-relevant pre-training (HateBERT, Twitter-RoBERTa) provides modest but '
     'measurable advantages over general-purpose encoders. Fine-tuning on labelled '
     'data is required for robust classification performance.')

# ═══════════════════════════════════════════════════════════════
# SECTION 2 — FINE-TUNING
# ═══════════════════════════════════════════════════════════════
heading(doc, '2.  Fine-Tuning BERT on Hate Speech Datasets  (20 Credits)', level=1)

heading(doc, '2.1  Approach', level=2)
body(doc,
     'We fine-tune BERT-base-uncased separately on two benchmark datasets — HateXPlain and '
     'TweetEval — for 3-class sequence classification (Normal / Offensive / Hate). A custom '
     'ManualAdamW optimizer is implemented to bypass a PyTorch 2.6 fused CUDA kernel '
     'incompatibility on the RTX 4060 GPU. Hyperparameter search covers learning rates in '
     '{1e-5, 2e-5, 3e-5} with batch size 8 and 3 training epochs per candidate. The best '
     'checkpoint (by validation Macro F1) is retained for each dataset.')

body(doc,
     'To address class imbalance (Hate speech is underrepresented in both datasets), class '
     'weights are computed via sklearn\'s compute_class_weight(\'balanced\') and applied '
     'through a weighted CrossEntropyLoss. Fine-tuned models are persisted as '
     'saved_model_bert_hx (HateXPlain) and saved_model_bert_te (TweetEval).')

heading(doc, '2.2  Results on Original Social-Media Test Sets', level=2)
body(doc,
     'Each fine-tuned model is first evaluated on its corresponding held-out test split '
     '(HateXPlain test set for BERT-HateXplain; TweetEval test set for BERT-TweetEval). '
     'Both models achieve solid performance on the social-media domain they were trained on:')

caption(doc, 'Table 2.  Fine-tuned BERT — Performance on Original Social-Media Domain Test Sets')
make_table(doc,
    headers=['Model', 'Test Set', 'Accuracy', 'Macro F1', 'F1 Normal', 'F1 Offensive', 'F1 Hate'],
    rows=[
        ['BERT-HateXplain', 'HateXPlain', '0.669', '0.662', '0.717', '0.536', '0.733'],
        ['BERT-TweetEval',  'TweetEval',  '0.579', '0.536', '0.536', '0.432', '0.639'],
    ],
    col_widths_cm=[3.5, 2.8, 2.0, 2.0, 2.0, 2.5, 2.0]
)

body(doc,
     'Both models recognise all three classes effectively in the social-media domain. '
     'BERT-HateXplain achieves F1 Hate = 0.733 on the HateXPlain test set, and '
     'BERT-TweetEval achieves F1 Hate = 0.639 on TweetEval. The confusion matrices below '
     'confirm that predictions are well distributed across all three classes.')

# Insert confusion matrix images side by side (original domain)
insert_two_images(doc,
    os.path.join(IMG_DIR, 'bert_finetune_hatexplain_cm.png'),
    'Figure 1. BERT-HateXplain on HateXPlain test set',
    os.path.join(IMG_DIR, 'bert_finetune_tweeteval_cm.png'),
    'Figure 2. BERT-TweetEval on TweetEval test set',
    w=2.9)

doc.add_paragraph()  # spacer

heading(doc, '2.3  Cross-Domain Evaluation: Course Review Test Set', level=2)
body(doc,
     'To assess real-world applicability, both fine-tuned models are additionally evaluated '
     'on a custom course-review test set comprising 286 samples collected from university '
     'course evaluations (50% Normal / 25% Offensive / 25% Hate). This dataset covers '
     'institutional offensive language — negative course reviews, personal attacks on '
     'instructors — which is stylistically distinct from the social-media content the '
     'models were trained on.')

caption(doc, 'Table 3.  Fine-tuned BERT — Performance on Course Review Test Set (n = 286)')
make_table(doc,
    headers=['Model', 'Test Set', 'Accuracy', 'Macro F1', 'F1 Normal', 'F1 Offensive', 'F1 Hate'],
    rows=[
        ['BERT-HateXplain', 'Course (OOD)', '0.535', '0.295', '0.714', '0.172', '0.000'],
        ['BERT-TweetEval',  'Course (OOD)', '0.514', '0.252', '0.714', '0.043', '0.000'],
    ],
    col_widths_cm=[3.5, 2.8, 2.0, 2.0, 2.0, 2.5, 2.0]
)

body(doc,
     'Performance collapses on the course domain. F1 Hate drops to 0.000 for both models — '
     'neither model successfully identifies a single course-domain Hate sample. F1 Offensive '
     'also degrades sharply (0.536 → 0.172 for BERT-HateXplain; 0.432 → 0.043 for '
     'BERT-TweetEval). The models default to predicting Normal for nearly all inputs, '
     'which is confirmed by the confusion matrices below.')

insert_two_images(doc,
    os.path.join(IMG_DIR, 'course_test_BERT-HateXplain_3-class_cm.png'),
    'Figure 3. BERT-HateXplain on course test (OOD)',
    os.path.join(IMG_DIR, 'course_test_BERT-TweetEval__3-class_cm.png'),
    'Figure 4. BERT-TweetEval on course test (OOD)',
    w=2.9)

doc.add_paragraph()

body(doc,
     'This reveals a severe domain shift problem: social-media hate speech classifiers '
     'trained on racial and religious slurs do not recognise course-domain offensive language '
     'such as "this lecturer is an aggressive bully" or "fucking trash course". The vocabulary, '
     'targets, and linguistic register are fundamentally different.')

# Side-by-side comparison
caption(doc, 'Table 4.  Domain Shift Summary — Original vs. Course Domain Performance')
make_table(doc,
    headers=['Model', 'Original Domain\nMacro F1', 'Course Domain\nMacro F1',
             'F1 Hate (orig.)', 'F1 Hate (course)'],
    rows=[
        ['BERT-HateXplain', '0.662', '0.295', '0.733', '0.000'],
        ['BERT-TweetEval',  '0.536', '0.252', '0.639', '0.000'],
    ],
    col_widths_cm=[3.8, 3.0, 3.0, 2.5, 2.5]
)

heading(doc, '2.4  Domain Adaptation Fine-Tuning', level=2)
body(doc,
     'To resolve the domain shift, domain adaptation fine-tuning is applied starting from '
     'saved_model_bert_hx — the model with the highest baseline performance on the course '
     'domain. Fine-tuning continues on 251 course-domain training samples using a small '
     'learning rate (2e-5) to avoid catastrophic forgetting of the original social-media '
     'knowledge. Class weights (Normal=0.450 / Offensive=1.946 / Hate=3.803) address the '
     'imbalanced course-domain distribution (74% / 17% / 9%), and training runs for 5 epochs.')

caption(doc, 'Table 5.  BERT-HateXplain Before and After Domain Adaptation (Course Test Set, n = 286)')
make_table(doc,
    headers=['Stage', 'Accuracy', 'Macro F1', 'F1 Normal', 'F1 Offensive', 'F1 Hate'],
    rows=[
        ['Before Adaptation (social-media model)', '0.535', '0.295', '0.714', '0.172', '0.000'],
        ['After Adaptation  (domain-adapted)',     '0.888', '0.863', '0.948', '0.765', '0.877'],
    ],
    col_widths_cm=[5.2, 2.0, 2.0, 2.2, 2.6, 2.0]
)

# Domain adaptation confusion matrix
p_img = doc.add_paragraph()
p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.isfile(os.path.join(IMG_DIR, 'course_domain_adapt_cm.png')):
    run = p_img.add_run()
    run.add_picture(os.path.join(IMG_DIR, 'course_domain_adapt_cm.png'), width=In_(3.1))
cap_p = doc.add_paragraph()
cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = cap_p.add_run('Figure 5. BERT-HateXplain after domain adaptation — course test set')
fmt(cap_run, italic=True, size=9, color=(0x60, 0x60, 0x60))

doc.add_paragraph()

body(doc,
     'Domain adaptation with only 251 course-domain samples raised Macro F1 from 0.295 to '
     '0.863 and lifted F1 Hate from 0.000 to 0.877, demonstrating the effectiveness of '
     'continual learning with a restrained learning rate to prevent catastrophic forgetting '
     'of the original social-media knowledge.')

# ═══════════════════════════════════════════════════════════════
# SECTION 3 — CASCADED (FAILED)
# ═══════════════════════════════════════════════════════════════
heading(doc, '3.  Cascaded Classification — Extended Method Attempt  (30 Credits)', level=1)

heading(doc, '3.1  Motivation', level=2)
body(doc,
     'Given the class imbalance inherent in hate speech data (Hate is rare; Normal is '
     'dominant), a cascaded two-stage architecture was designed as a method extension. '
     'The hypothesis was that decomposing the 3-class problem into two simpler binary '
     'decisions would reduce confusion between adjacent classes and improve recall on '
     'minority classes:')
bullet(doc, 'Stage 1 — Binary classifier: Normal vs. Not-Normal')
bullet(doc, 'Stage 2 — Binary classifier: Offensive vs. Hate '
            '(applied only to samples predicted as Not-Normal by Stage 1)')

heading(doc, '3.2  Implementation', level=2)
body(doc,
     'Both stages use BERT-base-uncased fine-tuned with ManualAdamW, weighted '
     'CrossEntropyLoss, and learning-rate search over {1e-5, 2e-5, 3e-5}. '
     'Course-domain data was mixed into training sets to attempt cross-domain '
     'generalisation. Stage models are saved as saved_model_bert_hx_stage1 / stage2 '
     '(HateXPlain) and saved_model_bert_te_stage1 / stage2 (TweetEval).')

heading(doc, '3.3  Results', level=2)
caption(doc, 'Table 6.  Cascaded vs. Direct Fine-Tuning — Macro F1 Comparison')
make_table(doc,
    headers=['Method', 'Dataset', 'Macro F1', 'Change vs. Direct'],
    rows=[
        ['Direct fine-tuning',         'HateXPlain', '0.731', '—'],
        ['Cascaded + course data mix',  'HateXPlain', '0.727', '−0.004'],
        ['Direct fine-tuning',         'TweetEval',  '0.536', '—'],
        ['Cascaded + course data mix',  'TweetEval',  '0.404', '−0.132'],
    ],
    col_widths_cm=[5.5, 3.0, 2.5, 3.8]
)

heading(doc, '3.4  Analysis and Conclusion', level=2)
body(doc,
     'The cascaded approach failed to improve performance and substantially degraded '
     'TweetEval Macro F1 from 0.536 to 0.404 (−13.2 points). Two principal failure '
     'mechanisms were identified:')

add_inline_bold(doc, [
    ('(1)  Error propagation. ', True),
    ('Misclassifications in Stage 1 cascade irreversibly into Stage 2; '
     'any Normal sample wrongly routed to Stage 2 cannot be recovered, '
     'compounding errors across stages.', False)
])

add_inline_bold(doc, [
    ('(2)  Domain mismatch from data mixing. ', True),
    ('Course-domain Normal text (structured academic descriptions) is '
     'linguistically distinct from Twitter Normal text. Mixing them '
     'confused Stage 1\'s decision boundary, causing many Normal course '
     'reviews to be incorrectly routed to Stage 2.', False)
])

body(doc,
     'These findings demonstrate that architectural complexity does not substitute for '
     'domain alignment. The cascaded approach was abandoned in favour of domain adaptation, '
     'which proved substantially more effective with far less design overhead.')

# ═══════════════════════════════════════════════════════════════
# SECTION 4 — SUMMARY
# ═══════════════════════════════════════════════════════════════
heading(doc, '4.  Summary and Key Findings', level=1)

caption(doc, 'Table 7.  Overall BERT Model Progression — Course Domain Test Set (n = 286)')
make_table(doc,
    headers=['Model / Stage', 'Accuracy', 'Macro F1', 'F1 Normal', 'F1 Offensive', 'F1 Hate'],
    rows=[
        ['BERT-base-uncased (zero-shot)',   '~0.40', '~0.37',  '—',    '—',    '—'],
        ['BERT-HateXplain (fine-tuned)',    '0.535', '0.295', '0.714', '0.172', '0.000'],
        ['BERT-TweetEval  (fine-tuned)',    '0.514', '0.252', '0.714', '0.043', '0.000'],
        ['BERT-HateXplain + Cascaded†',    '—',     '0.727*', '—',    '—',     '—'],
        ['BERT-HateXplain + Domain Adapt', '0.888', '0.863', '0.948', '0.765', '0.877'],
    ],
    col_widths_cm=[5.0, 2.0, 2.0, 2.2, 2.8, 2.0]
)

note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(2)
note_p.paragraph_format.space_after  = Pt(6)
run = note_p.add_run(
    '† Cascaded result on original HateXPlain test set (not course domain). '
    '* Lower on course domain due to domain shift.'
)
fmt(run, italic=True, size=9, color=(0x60, 0x60, 0x60))

body(doc,
     'The experimental progression from zero-shot (Macro F1 ≈ 0.37) through direct '
     'fine-tuning (0.295 on course domain) to domain adaptation (0.863 on course domain) '
     'demonstrates three key findings:')

bullet(doc,
       'Domain shift is the primary bottleneck. Models trained on social-media hate speech '
       'generalise poorly to course-evaluation offensive language regardless of architecture.')
bullet(doc,
       'Domain adaptation is highly sample-efficient. Only 251 course-domain training '
       'samples were sufficient to raise F1 Hate from 0.000 to 0.877.')
bullet(doc,
       'Architectural complexity does not compensate for domain mismatch. The cascaded '
       'two-stage design introduced error propagation without improving cross-domain '
       'generalisation.')

# ═══════════════════════════════════════════════════════════════
# SECTION 5 — COMMAND LINE TESTING
# ═══════════════════════════════════════════════════════════════
heading(doc, '5.  Command Line Interface Testing  (5 Credits)', level=1)

heading(doc, '5.1  Overview', level=2)
body(doc,
     'A command-line interface (CLI) tool, predict.py, is provided to allow interactive '
     'testing of the hate speech detection system. The tool accepts either a single text '
     'string or a plain-text file as input, runs inference through the selected model, '
     'and outputs a predicted label (Normal / Offensive / Hate) together with a confidence '
     'score for each input sample.')

heading(doc, '5.2  Design', level=2)
body(doc,
     'The CLI is built with Python\'s argparse module and wraps the full inference pipeline '
     'in a lightweight, dependency-free script. Key design decisions:')

add_inline_bold(doc, [
    ('Default model. ', True),
    ('The tool automatically loads saved_model_roberta_hx_course — the best-performing '
     'model (Macro F1 = 0.936 on the course test set) — when no model is explicitly '
     'specified. Any saved model directory can be selected via the --model flag, '
     'enabling comparison across all trained checkpoints.', False)
])

add_inline_bold(doc, [
    ('Dual input modes. ', True),
    ('--text accepts a single quoted string for quick testing; --file reads a '
     'plain-text file where each line is treated as one input sample, enabling '
     'batch evaluation without code changes.', False)
])

add_inline_bold(doc, [
    ('Confidence output. ', True),
    ('The predicted class label and its softmax probability (confidence score) are '
     'printed for every input. The --verbose flag additionally prints the full '
     'per-class probability distribution (Normal / Offensive / Hate), making it '
     'easy to inspect borderline predictions.', False)
])

add_inline_bold(doc, [
    ('Device-agnostic. ', True),
    ('The tool automatically selects CUDA if available, falling back to CPU, so '
     'it runs identically on both the development GPU and a standard laptop.', False)
])

heading(doc, '5.3  Usage', level=2)
body(doc, 'Basic usage examples:')

# Code-style table for commands
code_rows = [
    ['Single text input',
     'python predict.py --text "This course is complete garbage."'],
    ['Batch file input',
     'python predict.py --file reviews.txt'],
    ['Verbose (all class probs)',
     'python predict.py --text "The tutor was patient." --verbose'],
    ['Select a different model',
     'python predict.py --text "..." --model saved_model_bert_hx_course'],
]
caption(doc, 'Table 8.  predict.py Command Examples')
make_table(doc,
    headers=['Mode', 'Command'],
    rows=code_rows,
    col_widths_cm=[4.0, 10.8]
)

heading(doc, '5.4  Sample Output', level=2)
body(doc, 'Example output for a single text input with --verbose:')

# Monospace code block via a bordered paragraph
code_lines = [
    'Loading model from: saved_model_roberta_hx_course',
    'Device: cuda',
    '',
    'Text:       This lecturer is an aggressive bully who poisons every class.',
    'Prediction: Hate  (confidence: 0.9821)',
    '  Normal    : 0.0089',
    '  Offensive : 0.0090',
    '  Hate      : 0.9821',
]
for line in code_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run(line if line else ' ')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    # light grey background on code block
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)

doc.add_paragraph()  # spacer

heading(doc, '5.5  Workflow', level=2)
body(doc,
     'The end-to-end inference pipeline executed by predict.py proceeds as follows:')

steps = [
    ('Input parsing. ',
     'argparse reads --text or --file. If a file path is provided, each non-empty '
     'line is loaded as an independent input sample.'),
    ('Model loading. ',
     'AutoTokenizer and AutoModelForSequenceClassification are instantiated from '
     'the specified model directory (HuggingFace format). The model is moved to '
     'GPU if available and set to eval() mode.'),
    ('Tokenisation. ',
     'Each input string is tokenised with padding to max_length=128 and truncation. '
     'The resulting input_ids and attention_mask tensors are transferred to the '
     'target device.'),
    ('Inference. ',
     'A forward pass is performed under torch.no_grad(). The output logits are '
     'converted to class probabilities via softmax.'),
    ('Output. ',
     'The argmax of the probability vector determines the predicted label. '
     'The corresponding probability is reported as the confidence score. '
     'With --verbose, all three class probabilities are displayed.'),
]
for title, desc in steps:
    add_inline_bold(doc, [(f'{steps.index((title,desc))+1}.  {title}', True), (desc, False)], space_after=4)

# ── Save ──────────────────────────────────────────────────────
OUT = r'C:\Users\chenf\Desktop\6713\group\dataset\6713-ass2-dataset\model_code\v2\bert_report.docx'
doc.save(OUT)
print(f'Saved: {OUT}')
